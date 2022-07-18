# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

#%%
import docx
from translate import Translator
translator= Translator(from_lang="zh",to_lang="en")

from xpinyin import Pinyin


#%%

doc = docx.Document(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")

#%%

def name_trans(name):
    
    p = Pinyin()
    result1 = p.get_pinyin(name)   #convert to english
    #print(result1)
    s = result1.split('-')
    translation = ''.join(s[1:]).capitalize() + ' ' + ''.join(s[0]).capitalize()
   # print(translation)
    return translation
    
#%%
for i in range(7,7+20):
    print(doc.tables[5].cell(i,1).text)
    zh_name = doc.tables[5].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[5].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[5].cell(i,1).text)

#%%

for i in range(36,36+23):
    print(doc.tables[5].cell(i,1).text)
    zh_name = doc.tables[5].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[5].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[5].cell(i,1).text)

#%%
for i in range(67,67+25):
    print(doc.tables[5].cell(i,1).text)
    zh_name = doc.tables[5].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[5].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[5].cell(i,1).text)

#%%
for i in range(92,92+8):
    print(doc.tables[5].cell(i,1).text)
    zh_name = doc.tables[5].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[5].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[5].cell(i,1).text)

#%%
for i in range(109,109+9):
    print(doc.tables[5].cell(i,1).text)
    zh_name = doc.tables[5].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[5].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[5].cell(i,1).text)

#%%
for i in range(7,7+24):
    print(doc.tables[6].cell(i,1).text)
    zh_name = doc.tables[6].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[6].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[6].cell(i,1).text)
#%%
for i in range(7+23+10,7+23+10+26+1):
    print(doc.tables[6].cell(i,1).text)
    zh_name = doc.tables[6].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[6].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[6].cell(i,1).text)
#%%
for i in range(67+8,67+8+29):
    print(doc.tables[6].cell(i,1).text)
    zh_name = doc.tables[6].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[6].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[6].cell(i,1).text)
#%%
for i in range(113,123):
    print(doc.tables[6].cell(i,1).text)
    zh_name = doc.tables[6].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[6].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[6].cell(i,1).text)
#%%
for i in range(7,28):
    print(doc.tables[7].cell(i,1).text)
    zh_name = doc.tables[7].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[7].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[7].cell(i,1).text)
#%%
for i in range(37,47):
    print(doc.tables[7].cell(i,1).text)
    zh_name = doc.tables[7].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[7].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[7].cell(i,1).text)

#%%
for i in range(7,18):
    print(doc.tables[8].cell(i,1).text)
    zh_name = doc.tables[8].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[8].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[8].cell(i,1).text)
#%%
for i in range(6,18):
    print(doc.tables[9].cell(i,1).text)
    zh_name = doc.tables[9].cell(i,1).text
    trans = zh_name + ' ' + name_trans(zh_name)
    doc.tables[9].cell(i,1).text = trans
    doc.save(r"C:\Users\Shenghui\Desktop\cocm\GuestInfo_Pack_26.07.22_V3.docx")
    print(doc.tables[9].cell(i,1).text)